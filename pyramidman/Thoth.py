
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


import pandas as pd
import datetime as dt
import os
import gensim

from typing import List
# Creare new Word Document


class Papyrus():
    """Class to generate a report for the meetings assisted by pyradmidman.
    """

    def __init__(self, title: str = "pyramidman presentation", report_type: str = "meeting",
                 date: str = None, attendants: List[str] = None):
        # Main report information
        self.title = title
        self.report_type = report_type
        if date is None:
            self.date = str(dt.date.today())
        else:
            self.date = date

        self.attendants = attendants

        # Initialize some data structures
        self.document = None

        # Hardcoded param:
        self.logo_filepath = "../img/pyramidman_logo2.png"

    def summarize(self,text):
        return gensim.summarization.summarizer.summarize(text, word_count = 80)

    def set_transcription(self, transcription):
        """Sets the transcription of the minutes and its summary
        """
        self.transcription = transcription
        self.transcription_summary = self.summarize(transcription)

    def add_table_from_df(self, df, style = "Colorful Grid Accent 2"):
        """It creates and adds a table to the document from a pandas table
        """
        nrows, ncols = df.shape
        columns = df.columns.values
        table = self.document.add_table(rows=nrows+1, cols=ncols, style = style)

        header_cells = table.rows[0].cells
        i = 0
        for col in columns:
            header_cells[i].text = col
            i += 1

        for i in range(nrows):
            row_cells = table.rows[i+1].cells
            for j in range(ncols):
                row_cells[j].text = str(df.iloc[i][columns[j]])

    def add_first_page(self):

        ######################## FIRST PAGE #####################
        # Add pyramidman logo
        self.document.add_picture(self.logo_filepath, width=Inches(3.0))
        self.document.add_heading('Meeting:  ' + self.title, level=0)

        p = self.document.add_paragraph()

        # Now we add different texts of the first paragraph.
        text = 'This document contains a summary and transcription of the meeting '
        p.add_run(text)
        text = self.title + " "
        p.add_run(text).bold = True

        text = "\nMeeting Date: "
        p.add_run(text).bold = True
        text = str(self.date) + " "
        p.add_run(text)

        text = "\nAttendants: "
        p.add_run(text).bold = True
        text = ""
        
        if self.attendants is not None:
            for i in range(len(self.attendants)):
                text += self.attendants[i]
                if i < len(self.attendants)-1:
                    text += ", "

        p.add_run(text)

        self.document.add_heading('Summary ', level=1)
        p = self.document.add_paragraph()
        text = self.transcription_summary
        p.add_run(text)

        self.document.add_picture(self.word_cloud_image_path, width=Inches(3.0))
        last_paragraph = self.document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.document.add_page_break()

    def add_second_page(self):
        self.document.add_heading('Transcript ', level=1)
        p = self.document.add_paragraph()
        text = self.transcription
        p.add_run(text)


    def create_document(self, filepath='presentation.docx'):
        # This function will create the document, assuming all the data is loaded

        # Create word document
        document = Document()
        self.document = document
        self.add_first_page()
        self.add_second_page()

        # Save document in the end
        document.save(filepath)
        os.system("chmod 777 " + filepath)
